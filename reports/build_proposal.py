"""Generate a 2-page research proposal as a Word document.

Output: reports/proposal_2page.docx
"""
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, Cm, RGBColor


# ── Helpers ───────────────────────────────────────────────────────────────────
def set_cell_shading(cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def set_cell_borders(cell, color: str = "808080", size: str = "4") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), size)
        b.set(qn("w:color"), color)
        tc_borders.append(b)
    tc_pr.append(tc_borders)


def add_run(paragraph, text: str, *, bold=False, italic=False, size=10, color=None):
    run = paragraph.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor(*color)
    return run


def add_heading(doc, text: str, *, size=11, color=(31, 78, 121), space_before=4, space_after=2):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.keep_with_next = True
    add_run(p, text, bold=True, size=size, color=color)
    return p


def add_body(doc, text: str, *, size=10, space_after=2, justify=True, first_line_indent=None):
    p = doc.add_paragraph()
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15
    if first_line_indent:
        p.paragraph_format.first_line_indent = Cm(first_line_indent)
    # Allow bold via **...** segments
    parts = text.split("**")
    for i, seg in enumerate(parts):
        add_run(p, seg, bold=(i % 2 == 1), size=size)
    return p


def add_bullet(doc, text: str, *, size=10):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.line_spacing = 1.1
    parts = text.split("**")
    for i, seg in enumerate(parts):
        add_run(p, seg, bold=(i % 2 == 1), size=size)
    return p


def add_table(doc, rows, *, col_widths_cm=None, header_fill="1F4E79", header_color=(255, 255, 255), size=9):
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.autofit = False
    if col_widths_cm:
        for i, w in enumerate(col_widths_cm):
            for cell in table.columns[i].cells:
                cell.width = Cm(w)
    for r, row in enumerate(rows):
        for c, cell_text in enumerate(row):
            cell = table.cell(r, c)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_borders(cell, color="A6A6A6", size="4")
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            if r == 0:
                set_cell_shading(cell, header_fill)
                add_run(p, cell_text, bold=True, size=size, color=header_color)
            else:
                parts = cell_text.split("**")
                for i, seg in enumerate(parts):
                    add_run(p, seg, bold=(i % 2 == 1), size=size)
    return table


# ── Build document ────────────────────────────────────────────────────────────
doc = Document()

# Compact margins to fit 2 pages
for section in doc.sections:
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.7)
    section.right_margin = Cm(1.7)

# Default style
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(10)
style.paragraph_format.space_after = Pt(2)
style.paragraph_format.line_spacing = 1.15

# ── Title block ───────────────────────────────────────────────────────────────
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_after = Pt(2)
add_run(
    title,
    "Choosing the Right Simulated Learner: A Framework-Aligned Comparison "
    "of Teachable Agent and Student Simulation Protocols",
    bold=True, size=14, color=(31, 78, 121),
)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.paragraph_format.space_after = Pt(6)
add_run(
    subtitle,
    "Research Proposal — Target venue: Computers & Education (Special Issue: GenAI Enhanced Learning, 2026)",
    italic=True, size=10, color=(89, 89, 89),
)

# ── 1. Problem & Significance ─────────────────────────────────────────────────
add_heading(doc, "1. Problem and Significance")
add_body(
    doc,
    "LLM-based simulated learners are increasingly used in educational technology research and design, "
    "yet two structurally distinct prompting paradigms — **Teachable Agent (TA)** (Blair et al., 2007) "
    "and **Student Simulation (SS)** (Koedinger et al., 2015) — are routinely conflated under generic "
    "\"act-as-a-student\" prompts. We argue these protocols are **not interchangeable**: they target "
    "different educational purposes (learning-by-teaching vs. diagnostic/instructional testing) and "
    "therefore should be evaluated against different criteria. The proposed study provides the first "
    "**pre-registered, framework-aligned, multi-model comparative evaluation** of these protocols on "
    "matched public tutoring scenarios, yielding concrete design guidance for practitioners and a "
    "reusable benchmark for the GenAI-in-education community.",
)

# ── 2. Research Questions and Hypotheses ──────────────────────────────────────
add_heading(doc, "2. Research Questions and Pre-registered Hypotheses")
add_body(
    doc,
    "**RQ.** Do TA and SS protocols, implemented as LLM system prompts, produce structurally different "
    "simulated-learner evidence under matched tutoring scenarios, and is the difference stable across "
    "LLM families and datasets?",
)
add_bullet(doc, "**H1 (TA superiority):** C1 > {C2, C3, C4} on TA1–TA4 (Blair); Holm-corrected p < .0125, d > 0.5.")
add_bullet(doc, "**H2 (SS superiority):** C2 > {C3, C4} on SS1–SS5 (Koedinger); same thresholds.")
add_bullet(doc, "**H3 (Cross-model stability):** Direction of H1/H2 replicates in ≥ 3 of 4 LLM families.")
add_bullet(doc, "**H4 (Cross-dataset stability):** Effects replicate in ≥ 70 % of contrasts on Bridge.")
add_bullet(doc, "**H5 (Ablation separability):** Removing each structural element drops its target metric (8 ablations).")
add_bullet(doc, "**Equivalence (secondary):** C3 ≈ C4 on all 9 dimensions (TOST, bound d = 0.3).")

# ── 3. Experimental Design ────────────────────────────────────────────────────
add_heading(doc, "3. Experimental Design")
add_body(
    doc,
    "We adopt a fully crossed **scenario × condition × seed × model** design on the public "
    "**MathDial** corpus (Macina et al., 2023; CC-BY-4.0), with cross-dataset replication on **Bridge** "
    "(Wang et al., 2024). For every scenario, all conditions receive **identical tutor stimuli** — "
    "MathDial human-tutor turns are replayed verbatim — so any condition difference cannot be "
    "attributed to differential tutor behaviour. A near-transfer probe (same structure, different "
    "numbers) is appended as the final turn in every condition.",
)

add_table(
    doc,
    rows=[
        ["ID", "Condition", "Framework / Constraint", "Required structural elements"],
        ["C1", "**Teachable Agent**", "Blair et al. (2007)",
         "Ask questions • Explicit revision • Student stance • Independent transfer"],
        ["C2", "**Student Simulation**", "Koedinger et al. (2015)",
         "Target misconception • Verbalised reasoning • Gradual learning • Prior-knowledge boundary"],
        ["C3", "Generic Learner", "Baseline", "Single instruction: \"Act as a math student.\""],
        ["C4", "No-role Assistant", "Control", "Standard helpful-assistant prompt"],
    ],
    col_widths_cm=[1.0, 3.2, 3.8, 8.5],
)

add_body(
    doc,
    "**Sample size.** Main: 100 scenarios × 4 conditions × 3 seeds × 4 models = **4 800 dialogues**. "
    "Cross-dataset: 80 Bridge scenarios × 4 × 2 × 2 = **1 280 dialogues**. Ablations: 25 × 8 × 2 × 1 = "
    "**400 dialogues**. Human coding subset: **240 dialogues** (15 per condition × model). Scenarios are "
    "stratified by topic × error-type × difficulty (χ² independence tests, all p > .17).",
    space_after=3,
)

add_table(
    doc,
    rows=[
        ["Model family", "Model ID", "Role"],
        ["OpenAI", "gpt-4o-2024-11-20", "Primary"],
        ["Anthropic", "claude-sonnet-4-5-20250929", "Cross-family replication"],
        ["Google", "gemini-2.0-flash-001 (via OpenRouter)", "Cross-RLHF replication"],
        ["Meta (open)", "Llama-3.1-70B-Instruct", "Reproducibility anchor"],
    ],
    col_widths_cm=[3.0, 7.0, 6.5],
)

add_body(
    doc,
    "**Frozen generation parameters.** Temperature 0.7; max_tokens 600/turn; seeds [17, 42, 91]; "
    "prompts v1.0 (commit c2c41e6, read-only). Six pre-registered exclusion rules E1–E6 (empty turn, "
    "off-topic, refusal, overflow, missing transfer, role leakage) govern dialogue validity; rates are "
    "reported per condition × model.",
)

# ── 4. Measurement ────────────────────────────────────────────────────────────
add_heading(doc, "4. Measurement: Triangulating Human, Automatic and LLM-Judge Evidence")
add_bullet(
    doc,
    "**Human coding (primary):** 2 trained coders, 100 % overlap on 240 dialogues, blind to "
    "condition/model. 9 framework-aligned dimensions (TA1–TA4, SS1–SS5; 1–5 anchored scale). "
    "Gating reliability **ICC(2,k) ≥ 0.75**; arbitration by faculty third coder if |Δ| ≥ 2.",
)
add_bullet(
    doc,
    "**Automatic trace metrics (secondary):** 10 behaviour metrics computed on **all 6 480 valid "
    "dialogues** (question-asking rate, target-error preservation, feedback uptake, transfer attempt, "
    "premature correctness, role drift, etc.); validated against human ratings (Spearman ρ).",
)
add_bullet(
    doc,
    "**LLM-as-judge (generalisation):** different-family judge (e.g., Claude Opus 4 or GPT-5), 3-pass "
    "median; calibrated against the 240 human ratings — dimension included **only if ICC(judge, human) ≥ 0.65**; "
    "family-bias check (Δ ≥ 0.5 flagged).",
)

# ── 5. Statistical Analysis Plan ──────────────────────────────────────────────
add_heading(doc, "5. Statistical Analysis Plan")
add_body(
    doc,
    "**Primary model.** Linear mixed-effects (lme4 via pymer4): "
    "`rating ~ condition * model_family + (1|scenario_id) + (1|coder_id)`, REML estimation, "
    "deviation coding with C4 as reference. **Multiple comparisons:** Holm-Bonferroni over 4 primary "
    "contrast families (α = .0125); Benjamini-Hochberg FDR in supplement. **Effect sizes:** Cohen's d "
    "with bootstrap 95 % CI (5 000 iterations). **Equivalence:** TOST with bound d = 0.3 for C3-vs-C4. "
    "**Ablations:** paired Wilcoxon vs. full protocol per target metric. **Power:** MDE d ≈ 0.32 at "
    "80 % power for n = 100 scenario-matched pairs; preliminary trace differences imply d > 1.0.",
)

# ── 6. Robustness, Ethics, Open Science ───────────────────────────────────────
add_heading(doc, "6. Robustness, Ethics, and Open Science")
add_bullet(
    doc,
    "**Robustness grid:** condition × {topic, error_type, difficulty, model, seed, dataset} splits; "
    "direction-consistency reported as primary robustness statistic.",
)
add_bullet(
    doc,
    "**Pre-registration:** analysis plan with H1–H5, exclusions and stopping rule committed to public "
    "GitHub before main generation (commit 27a554e, 2026-05-15 UTC).",
)
add_bullet(
    doc,
    "**Ethics & reproducibility:** IRB-exempt (public corpora only, no human subjects); all code, "
    "prompts, 4 800 dialogues, ratings and a Datasheet released under MIT/CC-BY-4.0 on GitHub + Zenodo.",
)

# ── 7. Timeline and Deliverables ──────────────────────────────────────────────
add_heading(doc, "7. Timeline and Deliverables")
add_table(
    doc,
    rows=[
        ["Phase", "Weeks", "Output"],
        ["P1–P2  Scenario bank + prompt freeze", "1–2", "Stratified bank (100 + 50); frozen prompts v1.0"],
        ["P3  Pilot generation + codebook v1.1", "3", "400 pilot dialogues; reliability training set"],
        ["P4  Main generation (4 models)", "4–5", "4 800 dialogues + automatic trace metrics"],
        ["P5  Bridge replication + ablations", "6", "1 280 + 400 dialogues"],
        ["P6  Human coding (240, ICC gate)", "7–9", "Coder ratings; ICC ≥ 0.75 confirmed"],
        ["P7  LLM-judge calibration + scoring", "10", "Judge ratings for n = 4 800"],
        ["P8  Analysis + writing", "11–15", "Mixed-effects, TOST, ablations; manuscript draft"],
        ["P9  Submission to Computers & Education", "16", "arXiv preprint + journal submission"],
    ],
    col_widths_cm=[7.5, 2.5, 6.5],
)

add_body(
    doc,
    "**Expected contribution.** (i) The first multi-model, multi-dataset, framework-aligned benchmark "
    "for LLM simulated learners; (ii) a use-case decision tree mapping protocol → educational purpose; "
    "(iii) an open-source replication package establishing methodological standards for future "
    "GenAI-in-education evaluation studies.",
    space_after=0,
)

out = Path(__file__).parent / "proposal_2page.docx"
doc.save(out)
print(f"Saved: {out}")
